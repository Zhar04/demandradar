"""Сборка docs/PLAYBOOK.md -> docs/PLAYBOOK.docx (только опенсорс).

Путь 1: pandoc, если установлен (лучшее качество: таблицы, стили).
Путь 2: python-docx (pip install python-docx) — заголовки/абзацы/списки,
        таблицы переносятся упрощённо построчно.

Запуск: .venv\\Scripts\\python scripts/build_playbook_docx.py
"""

from __future__ import annotations

import re
import shutil
import subprocess
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "PLAYBOOK.md"
TARGET = ROOT / "docs" / "PLAYBOOK.docx"


def build_with_pandoc() -> bool:
    pandoc = shutil.which("pandoc")
    if pandoc is None:
        return False
    subprocess.run(
        [pandoc, str(SOURCE), "-o", str(TARGET), "--from", "gfm", "--toc"],
        check=True,
    )
    return True


def build_with_python_docx() -> None:
    try:
        import docx  # python-docx
    except ImportError:
        sys.exit(
            "Ни pandoc, ни python-docx не найдены.\n"
            "Вариант 1: choco install pandoc (или https://pandoc.org/installing)\n"
            "Вариант 2: .venv\\Scripts\\pip install python-docx"
        )

    document = docx.Document()
    bold_re = re.compile(r"\*\*(.+?)\*\*")

    def add_paragraph(text: str, style: str | None = None):
        paragraph = document.add_paragraph(style=style)
        pos = 0
        for match in bold_re.finditer(text):
            if match.start() > pos:
                paragraph.add_run(text[pos:match.start()])
            paragraph.add_run(match.group(1)).bold = True
            pos = match.end()
        if pos < len(text):
            paragraph.add_run(text[pos:])

    for raw_line in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 4)
            document.add_heading(stripped.lstrip("#").strip(), level=level)
        elif stripped.startswith(">"):
            add_paragraph(stripped.lstrip("> ").strip(), style="Intense Quote")
        elif stripped.startswith(("- ", "* ")):
            add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            add_paragraph(re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
        elif stripped.startswith("|"):
            if set(stripped.replace("|", "").replace("-", "").replace(":", "").strip()) == set():
                continue  # разделитель шапки таблицы
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            add_paragraph(" — ".join(c for c in cells if c))
        else:
            add_paragraph(stripped)

    document.save(TARGET)


def main() -> None:
    if not SOURCE.exists():
        sys.exit(f"Не найден {SOURCE}")
    if build_with_pandoc():
        print(f"OK (pandoc): {TARGET}")
    else:
        build_with_python_docx()
        print(f"OK (python-docx): {TARGET}")


if __name__ == "__main__":
    main()
